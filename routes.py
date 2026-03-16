"""
4uPDF API Routes - Phase 1 MVP
FastAPI routes for core PDF tools: merge, split, compress, convert
"""

import os
import io
import uuid
import time
import zipfile
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
import fitz

router = APIRouter()

# In-memory job tracking
jobs_db = {}

def generate_job_id() -> str:
    return str(uuid.uuid4())[:8]


# ========== MERGE PDF ==========
@router.post("/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    """
    Merge multiple PDF files into one.
    Phase 1 Tool: /merge-pdf
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    if len(files) < 2:
        raise HTTPException(status_code=400, detail="At least 2 PDF files required")

    for upload_file in files:
        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="File has no filename")
        if not upload_file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail=f"Invalid file type: {upload_file.filename}. Only PDF files allowed")

    try:
        merged_doc = fitz.open()

        for upload_file in files:
            content = await upload_file.read()
            if len(content) == 0:
                raise HTTPException(status_code=400, detail=f"Empty file: {upload_file.filename}")
            if len(content) > 50 * 1024 * 1024:
                raise HTTPException(status_code=413, detail=f"File too large: {upload_file.filename}. Maximum 50MB per file")

            pdf_stream = io.BytesIO(content)
            try:
                src_doc = fitz.open(stream=pdf_stream, filetype="pdf")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Invalid PDF file: {upload_file.filename}")

            merged_doc.insert_pdf(src_doc)
            src_doc.close()

        output_stream = io.BytesIO()
        merged_doc.save(output_stream)
        merged_doc.close()
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=merged.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Merge failed: {str(e)}")


# ========== SPLIT PDF ==========
@router.post("/split")
async def split_pdf(
    file: UploadFile = File(...),
    mode: str = Form("pages"),  # "pages" or "ranges"
    ranges: Optional[str] = Form(None)  # e.g., "1-3,5,7-9"
):
    """
    Split PDF by page ranges or individual pages.
    Phase 1 Tool: /split-pdf
    """
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    if mode not in ["pages", "ranges"]:
        raise HTTPException(status_code=400, detail="Invalid mode. Must be 'pages' or 'ranges'")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        pdf_stream = io.BytesIO(content)
        try:
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid PDF file")
        total_pages = len(doc)

        if mode == "pages":
            output_files = []
            for i in range(total_pages):
                page_doc = fitz.open()
                page_doc.insert_pdf(doc, from_page=i, to_page=i)
                page_stream = io.BytesIO()
                page_doc.save(page_stream)
                page_doc.close()
                page_stream.seek(0)
                output_files.append((f"page_{i+1}.pdf", page_stream))

            doc.close()

            if len(output_files) == 1:
                return StreamingResponse(
                    output_files[0][1],
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"}
                )

            zip_stream = io.BytesIO()
            with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zf:
                for filename, file_stream in output_files:
                    zf.writestr(filename, file_stream.getvalue())
            zip_stream.seek(0)

            return StreamingResponse(
                zip_stream,
                media_type="application/zip",
                headers={"Content-Disposition": "attachment; filename=split_pages.zip"}
            )

        elif mode == "ranges" and ranges:
            range_parts = ranges.split(",")
            output_files = []

            for idx, part in enumerate(range_parts):
                part = part.strip()
                if "-" in part:
                    start, end = map(int, part.split("-"))
                    start -= 1
                    end -= 1
                else:
                    start = end = int(part) - 1

                if start < 0 or end >= total_pages or start > end:
                    raise HTTPException(status_code=400, detail=f"Invalid range: {part}")

                range_doc = fitz.open()
                range_doc.insert_pdf(doc, from_page=start, to_page=end)
                range_stream = io.BytesIO()
                range_doc.save(range_stream)
                range_doc.close()
                range_stream.seek(0)
                output_files.append((f"range_{idx+1}.pdf", range_stream))

            doc.close()

            if len(output_files) == 1:
                return StreamingResponse(
                    output_files[0][1],
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"}
                )

            zip_stream = io.BytesIO()
            with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zf:
                for filename, file_stream in output_files:
                    zf.writestr(filename, file_stream.getvalue())
            zip_stream.seek(0)

            return StreamingResponse(
                zip_stream,
                media_type="application/zip",
                headers={"Content-Disposition": "attachment; filename=split_ranges.zip"}
            )

        else:
            raise HTTPException(status_code=400, detail="Invalid mode or missing ranges")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Split failed: {str(e)}")


# ========== COMPRESS PDF ==========
@router.post("/compress")
async def compress_pdf(
    file: UploadFile = File(...),
    level: str = Form("medium")  # "low", "medium", "high"
):
    """
    Compress PDF file by reducing image quality.
    Phase 1 Tool: /compress-pdf
    """
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    quality_map = {
        "low": 50,
        "medium": 70,
        "high": 85
    }

    if level not in quality_map:
        raise HTTPException(status_code=400, detail="Invalid compression level. Must be 'low', 'medium', or 'high'")

    quality = quality_map[level]

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        pdf_stream = io.BytesIO(content)
        try:
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid PDF file")

        for page_num in range(len(doc)):
            page = doc[page_num]
            images = page.get_images()

            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                pix = fitz.Pixmap(image_bytes)
                if pix.alpha:
                    pix = fitz.Pixmap(fitz.csRGB, pix)

                compressed_bytes = pix.pil_tobytes(format="JPEG", quality=quality)

                page.insert_image(
                    page.rect,
                    stream=compressed_bytes,
                    keep_proportion=True,
                    overlay=False
                )

        output_stream = io.BytesIO()
        doc.save(output_stream, garbage=4, deflate=True, clean=True)
        doc.close()
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=compressed.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Compression failed: {str(e)}")


# ========== PDF TO WORD ==========
@router.post("/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """
    Convert PDF to DOCX.
    Phase 1 Tool: /pdf-to-word
    Note: Basic implementation extracts text. Future: preserve formatting.
    """
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        pdf_stream = io.BytesIO(content)
        try:
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid PDF file")

        text_content = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_content.append(page.get_text())

        doc.close()

        full_text = "\n\n".join(text_content)

        try:
            from docx import Document
            docx = Document()
            for para in full_text.split("\n\n"):
                if para.strip():
                    docx.add_paragraph(para.strip())

            output_stream = io.BytesIO()
            docx.save(output_stream)
            output_stream.seek(0)

            return StreamingResponse(
                output_stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=converted.docx"}
            )
        except ImportError:
            text_stream = io.BytesIO(full_text.encode("utf-8"))
            return StreamingResponse(
                text_stream,
                media_type="text/plain",
                headers={"Content-Disposition": "attachment; filename=converted.txt"}
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


# ========== WORD TO PDF ==========
@router.post("/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """
    Convert DOCX to PDF.
    Phase 1 Tool: /word-to-pdf
    Note: Requires python-docx and reportlab or comtypes (Windows only).
    """
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    if not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Invalid file type. Only DOCX/DOC files allowed")

    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        docx_stream = io.BytesIO(content)
        try:
            doc = Document(docx_stream)
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid DOCX file")

        output_stream = io.BytesIO()
        pdf = SimpleDocTemplate(output_stream, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles["Normal"]))
                story.append(Spacer(1, 12))

        pdf.build(story)
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=converted.pdf"}
        )

    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx or reportlab not installed")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


# ========== JPG TO PDF ==========
@router.post("/jpg-to-pdf")
async def jpg_to_pdf(files: List[UploadFile] = File(...)):
    """
    Convert JPG/PNG images to PDF.
    Phase 1 Tool: /jpg-to-pdf
    """
    if not files:
        raise HTTPException(status_code=400, detail="At least 1 image required")

    allowed_types = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}
    for upload_file in files:
        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="File has no filename")
        ext = Path(upload_file.filename).suffix.lower()
        if ext not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Invalid file type: {upload_file.filename}. Allowed: JPG, PNG, GIF, BMP, TIFF")

    try:
        doc = fitz.open()

        for upload_file in files:
            content = await upload_file.read()
            if len(content) == 0:
                raise HTTPException(status_code=400, detail=f"Empty file: {upload_file.filename}")
            if len(content) > 50 * 1024 * 1024:
                raise HTTPException(status_code=413, detail=f"File too large: {upload_file.filename}. Maximum 50MB per file")

            try:
                pix = fitz.Pixmap(content)
                page = doc.new_page(width=pix.width, height=pix.height)
                page.insert_image(page.rect, stream=content)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Invalid image file: {upload_file.filename}")

        output_stream = io.BytesIO()
        doc.save(output_stream)
        doc.close()
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=images.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


# ========== PDF TO JPG ==========
@router.post("/pdf-to-jpg")
async def pdf_to_jpg(
    file: UploadFile = File(...),
    dpi: int = Form(150),
    pages: Optional[str] = Form(None)  # e.g., "1-3,5"
):
    """
    Convert PDF pages to JPG images.
    Phase 1 Tool: /pdf-to-jpg
    """
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    if dpi < 72 or dpi > 600:
        raise HTTPException(status_code=400, detail="DPI must be between 72 and 600")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        pdf_stream = io.BytesIO(content)
        try:
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
        except Exception as e:
            raise HTTPException(status_code=400, detail="Invalid PDF file")
        total_pages = len(doc)

        page_indices = []
        if pages:
            for part in pages.split(","):
                part = part.strip()
                if "-" in part:
                    start, end = map(int, part.split("-"))
                    page_indices.extend(range(start - 1, end))
                else:
                    page_indices.append(int(part) - 1)
        else:
            page_indices = list(range(total_pages))

        output_files = []
        for idx in page_indices:
            if idx < 0 or idx >= total_pages:
                continue

            page = doc[idx]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.pil_tobytes(format="JPEG")
            output_files.append((f"page_{idx+1}.jpg", io.BytesIO(img_bytes)))

        doc.close()

        if not output_files:
            raise HTTPException(status_code=400, detail="No valid pages to convert")

        if len(output_files) == 1:
            return StreamingResponse(
                output_files[0][1],
                media_type="image/jpeg",
                headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"}
            )

        zip_stream = io.BytesIO()
        with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zf:
            for filename, file_stream in output_files:
                zf.writestr(filename, file_stream.getvalue())
        zip_stream.seek(0)

        return StreamingResponse(
            zip_stream,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=pdf_pages.zip"}
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


# ========== PHASE 2: ROTATE PDF ==========
@router.post("/rotate")
async def rotate_pdf(
    file: UploadFile = File(...),
    angle: int = Form(90),
    pages: Optional[str] = Form(None)
):
    """Rotate PDF pages by 90, 180, or 270 degrees."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")
    if angle not in [90, 180, 270]:
        raise HTTPException(status_code=400, detail="Angle must be 90, 180, or 270")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)

        page_indices = _parse_pages(pages, total) if pages else list(range(total))

        for idx in page_indices:
            if 0 <= idx < total:
                doc[idx].set_rotation((doc[idx].rotation + angle) % 360)

        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=rotated.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Rotation failed: {str(e)}")


# ========== PHASE 2: DELETE PAGES ==========
@router.post("/delete-pages")
async def delete_pages(
    file: UploadFile = File(...),
    pages: str = Form(...)
):
    """Delete specified pages from a PDF."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)

        delete_indices = sorted(set(_parse_pages(pages, total)), reverse=True)
        # Filter to valid indices only
        valid_indices = [i for i in delete_indices if 0 <= i < total]
        if not valid_indices:
            raise HTTPException(status_code=400, detail=f"No valid pages to delete. Document has {total} pages.")
        if len(valid_indices) >= total:
            raise HTTPException(status_code=400, detail="Cannot delete all pages")

        for idx in valid_indices:
            doc.delete_page(idx)

        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=pages_removed.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete pages failed: {str(e)}")


# ========== PHASE 2: EXTRACT PAGES ==========
@router.post("/extract-pages")
async def extract_pages(
    file: UploadFile = File(...),
    pages: str = Form(...)
):
    """Extract specified pages from a PDF into a new PDF."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)
        extract_indices = _parse_pages(pages, total)

        if not extract_indices:
            raise HTTPException(status_code=400, detail="No valid pages to extract")

        new_doc = fitz.open()
        for idx in extract_indices:
            if 0 <= idx < total:
                new_doc.insert_pdf(doc, from_page=idx, to_page=idx)

        doc.close()
        output = io.BytesIO()
        new_doc.save(output)
        new_doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=extracted.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extract pages failed: {str(e)}")


# ========== PHASE 2: WATERMARK PDF ==========
@router.post("/watermark")
async def watermark_pdf(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENTIAL"),
    opacity: float = Form(0.15),
    angle: int = Form(45),
    font_size: int = Form(60)
):
    """Add text watermark to every page of a PDF."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")
    if not text.strip():
        raise HTTPException(status_code=400, detail="Watermark text cannot be empty")
    if opacity < 0.01 or opacity > 1.0:
        raise HTTPException(status_code=400, detail="Opacity must be between 0.01 and 1.0")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")

        for page in doc:
            rect = page.rect
            cx, cy = rect.width / 2, rect.height / 2
            tw = fitz.TextWriter(page.rect)
            font = fitz.Font("helv")
            tw.append((cx - len(text) * font_size * 0.25, cy), text, font=font, fontsize=font_size)
            tw.write_text(page, opacity=opacity, morph=(fitz.Point(cx, cy), fitz.Matrix(angle)))

        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=watermarked.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Watermark failed: {str(e)}")


# ========== PHASE 2: PROTECT PDF (PASSWORD) ==========
@router.post("/protect")
async def protect_pdf(
    file: UploadFile = File(...),
    password: str = Form(...),
    owner_password: str = Form(""),
    permissions: str = Form("all")
):
    """Add password protection to a PDF."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")
    if not password or len(password) < 1:
        raise HTTPException(status_code=400, detail="Password is required")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")

        perm = fitz.PDF_PERM_ACCESSIBILITY
        if permissions == "all":
            perm = int(
                fitz.PDF_PERM_ACCESSIBILITY
                | fitz.PDF_PERM_PRINT
                | fitz.PDF_PERM_COPY
                | fitz.PDF_PERM_ANNOTATE
            )
        elif permissions == "print":
            perm = int(fitz.PDF_PERM_ACCESSIBILITY | fitz.PDF_PERM_PRINT)
        elif permissions == "none":
            perm = int(fitz.PDF_PERM_ACCESSIBILITY)

        encrypt_method = fitz.PDF_ENCRYPT_AES_256
        owner_pw = owner_password if owner_password else password

        output = io.BytesIO()
        doc.save(output, encryption=encrypt_method, user_pw=password, owner_pw=owner_pw, permissions=perm)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=protected.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Protection failed: {str(e)}")


# ========== PHASE 2: UNLOCK PDF ==========
@router.post("/unlock")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    """Remove password protection from a PDF."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")

        if doc.is_encrypted:
            if not doc.authenticate(password):
                raise HTTPException(status_code=403, detail="Invalid password")

        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=unlocked.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unlock failed: {str(e)}")


# ========== PHASE 3: SPLIT BY BARCODE ==========
@router.post("/split-barcode")
async def split_by_barcode(
    file: UploadFile = File(...),
    dpi: int = Form(200)
):
    """Split PDF at pages containing barcodes/QR codes. Each barcode starts a new section."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        from pyzbar import pyzbar
        from PIL import Image
    except ImportError:
        raise HTTPException(status_code=501, detail="pyzbar not installed on server")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 100MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)
        mat = fitz.Matrix(dpi / 72, dpi / 72)

        # Scan pages for barcodes
        splits = []  # list of (page_idx, barcode_data)
        for i in range(total):
            pix = doc[i].get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            barcodes = pyzbar.decode(img)
            if barcodes:
                splits.append((i, barcodes[0].data.decode("utf-8", errors="replace")))

        if not splits:
            raise HTTPException(status_code=400, detail="No barcodes found in any page")

        # Build sections: each barcode page starts a new section
        sections = []
        for idx, (page_idx, barcode_data) in enumerate(splits):
            end = splits[idx + 1][0] if idx + 1 < len(splits) else total
            sections.append({"start": page_idx, "end": end, "label": barcode_data})

        # Create split PDFs
        output_files = []
        for sec_idx, sec in enumerate(sections):
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=sec["start"], to_page=sec["end"] - 1)
            buf = io.BytesIO()
            new_doc.save(buf)
            new_doc.close()
            buf.seek(0)
            safe_label = "".join(c if c.isalnum() or c in "-_ " else "_" for c in sec["label"])[:50]
            fname = f"{sec_idx + 1}_{safe_label}.pdf"
            output_files.append((fname, buf))

        doc.close()

        if len(output_files) == 1:
            return StreamingResponse(output_files[0][1], media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"})

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, buf in output_files:
                zf.writestr(fname, buf.getvalue())
        zip_buf.seek(0)

        return StreamingResponse(zip_buf, media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=barcode_split.zip"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Barcode split failed: {str(e)}")


# ========== PHASE 3: AUTO-RENAME PDF ==========
@router.post("/auto-rename")
async def auto_rename(
    files: List[UploadFile] = File(...),
    region: str = Form("top"),
    pattern: Optional[str] = Form(None),
    dpi: int = Form(150)
):
    """
    Auto-rename PDFs based on OCR-detected text from first page.
    Region: top, bottom, full. Optional regex pattern to extract specific text.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    try:
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
        ocr = RapidOCR()
    except ImportError:
        raise HTTPException(status_code=501, detail="OCR engine not available")

    region_map = {
        "top": (0.0, 0.0, 1.0, 0.25),
        "bottom": (0.0, 0.75, 1.0, 1.0),
        "full": (0.0, 0.0, 1.0, 1.0),
    }
    crop = region_map.get(region, region_map["top"])

    try:
        import re
        output_files = []

        for upload_file in files:
            if not upload_file.filename or not upload_file.filename.lower().endswith('.pdf'):
                continue

            content = await upload_file.read()
            if len(content) == 0 or len(content) > 50 * 1024 * 1024:
                continue

            doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
            page = doc[0]
            pw, ph = page.rect.width, page.rect.height
            clip = fitz.Rect(pw * crop[0], ph * crop[1], pw * crop[2], ph * crop[3])
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat, clip=clip)

            import numpy as np
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
            result, _ = ocr(img)
            texts = [r[1] for r in result] if result else []
            full_text = " ".join(texts)

            if pattern:
                match = re.search(pattern, full_text, re.IGNORECASE)
                name_part = match.group(1) if match else full_text[:60]
            else:
                name_part = full_text[:60]

            safe_name = "".join(c if c.isalnum() or c in "-_ ." else "_" for c in name_part).strip("_")
            if not safe_name:
                safe_name = upload_file.filename.replace(".pdf", "")

            buf = io.BytesIO(content)
            output_files.append((f"{safe_name}.pdf", buf))
            doc.close()

        if not output_files:
            raise HTTPException(status_code=400, detail="No valid PDFs processed")

        if len(output_files) == 1:
            return StreamingResponse(output_files[0][1], media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"})

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, buf in output_files:
                buf.seek(0)
                zf.writestr(fname, buf.read())
        zip_buf.seek(0)

        return StreamingResponse(zip_buf, media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=renamed_pdfs.zip"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auto-rename failed: {str(e)}")


# ========== PHASE 3: DOCUMENT TYPE DETECTION ==========
@router.post("/detect-type")
async def detect_document_type(
    files: List[UploadFile] = File(...),
    dpi: int = Form(150)
):
    """
    Detect document type (invoice, contract, delivery note, etc.) using OCR + keyword matching.
    Returns classification for each uploaded PDF.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    try:
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
        ocr = RapidOCR()
    except ImportError:
        raise HTTPException(status_code=501, detail="OCR engine not available")

    DOCUMENT_TYPES = {
        "invoice": ["invoice", "invoice number", "total", "tax", "vat", "subtotal", "bill to", "amount due", "payment terms"],
        "order": ["order", "order number", "purchase order", "po number", "order confirmation", "sales order", "order date"],
        "contract": ["contract", "agreement", "parties", "clauses", "signatures", "terms and conditions", "effective date"],
        "delivery_note": ["delivery note", "delivery", "transport", "shipment", "packing list", "shipping", "consignment"],
        "receipt": ["receipt", "fiscal receipt", "cash register", "payment received", "transaction", "payment confirmation"],
        "report": ["report", "summary", "analysis", "statistics", "results", "findings", "conclusions"],
        "letter": ["dear", "regards", "sincerely", "to whom", "subject", "yours truly", "best regards"],
    }

    try:
        results = []
        for upload_file in files:
            if not upload_file.filename or not upload_file.filename.lower().endswith('.pdf'):
                continue

            content = await upload_file.read()
            if len(content) == 0 or len(content) > 50 * 1024 * 1024:
                results.append({"file": upload_file.filename, "type": "unknown", "confidence": 0, "keywords": []})
                continue

            doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
            # Read text from first 2 pages
            full_text = ""
            for page_idx in range(min(2, len(doc))):
                page = doc[page_idx]
                text = page.get_text()
                if text.strip():
                    full_text += text + " "
                else:
                    # OCR fallback
                    mat = fitz.Matrix(dpi / 72, dpi / 72)
                    pix = page.get_pixmap(matrix=mat)
                    import numpy as np
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
                    result, _ = ocr(img)
                    if result:
                        full_text += " ".join([r[1] for r in result]) + " "

            doc.close()
            text_lower = full_text.lower()

            # Score each type
            scores = {}
            matched_keywords = {}
            for doc_type, keywords in DOCUMENT_TYPES.items():
                found = [kw for kw in keywords if kw.lower() in text_lower]
                scores[doc_type] = len(found)
                matched_keywords[doc_type] = found

            best_type = max(scores, key=scores.get) if max(scores.values()) > 0 else "unknown"
            best_score = scores.get(best_type, 0)
            confidence = min(100, best_score * 25)

            results.append({
                "file": upload_file.filename,
                "type": best_type,
                "confidence": confidence,
                "keywords": matched_keywords.get(best_type, []),
                "page_count": len(fitz.open(stream=io.BytesIO(content), filetype="pdf")),
                "text_preview": full_text[:200].strip()
            })

        return {"results": results, "total": len(results)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")


# ========== PHASE 4: ARCHIVE PROCESSOR ==========
@router.post("/process-archive")
async def process_archive(
    files: List[UploadFile] = File(...),
    action: str = Form("detect-split"),
    pattern: Optional[str] = Form(None),
    dpi: int = Form(150)
):
    """
    Phase 4 automation: Process a batch of scanned PDFs.
    Actions:
    - detect-split: Detect document types and split by type into folders
    - rename: Auto-rename based on OCR content
    - classify: Just classify without modifying
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    if action not in ["detect-split", "rename", "classify"]:
        raise HTTPException(status_code=400, detail="Invalid action")

    try:
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
        import re
        ocr = RapidOCR()
    except ImportError:
        raise HTTPException(status_code=501, detail="OCR engine not available")

    DOCUMENT_TYPES = {
        "invoices": ["invoice", "invoice number", "total", "tax", "vat"],
        "orders": ["order", "order number", "purchase order", "po number", "sales order"],
        "contracts": ["contract", "agreement", "parties", "terms and conditions"],
        "delivery_notes": ["delivery note", "delivery", "shipment", "packing list"],
        "other": [],
    }

    try:
        classified = {}

        for upload_file in files:
            if not upload_file.filename or not upload_file.filename.lower().endswith('.pdf'):
                continue

            content = await upload_file.read()
            if len(content) == 0 or len(content) > 100 * 1024 * 1024:
                continue

            doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
            # Get text from first page
            page = doc[0]
            text = page.get_text()
            if not text.strip():
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = page.get_pixmap(matrix=mat)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
                result, _ = ocr(img)
                text = " ".join([r[1] for r in result]) if result else ""

            text_lower = text.lower()

            # Classify
            doc_type = "other"
            for dtype, keywords in DOCUMENT_TYPES.items():
                if any(kw in text_lower for kw in keywords):
                    doc_type = dtype
                    break

            # Determine filename
            if action == "rename" and pattern:
                match = re.search(pattern, text, re.IGNORECASE)
                fname = match.group(1) if match else upload_file.filename
                fname = "".join(c if c.isalnum() or c in "-_ ." else "_" for c in fname)[:60]
                if not fname.lower().endswith('.pdf'):
                    fname += '.pdf'
            else:
                fname = upload_file.filename

            doc.close()

            if doc_type not in classified:
                classified[doc_type] = []
            classified[doc_type].append((fname, io.BytesIO(content)))

        if not classified:
            raise HTTPException(status_code=400, detail="No valid PDFs processed")

        if action == "classify":
            summary = {dtype: [f[0] for f in files_list] for dtype, files_list in classified.items()}
            return {"classification": summary, "total": sum(len(v) for v in summary.values())}

        # Build ZIP with folders per type
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for doc_type, file_list in classified.items():
                for fname, buf in file_list:
                    buf.seek(0)
                    zf.writestr(f"{doc_type}/{fname}", buf.read())
        zip_buf.seek(0)

        return StreamingResponse(zip_buf, media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=processed_archive.zip"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Archive processing failed: {str(e)}")


# ========== PHASE 2: SIGN PDF ==========
@router.post("/sign-pdf")
async def sign_pdf(
    file: UploadFile = File(...),
    signature_type: str = Form("text"),
    signature_text: Optional[str] = Form(None),
    signature_image: Optional[UploadFile] = File(None),
    position: str = Form("bottom-right"),
    pages: Optional[str] = Form(None),
    font_size: int = Form(12),
    opacity: float = Form(1.0)
):
    """Add signature to PDF pages (text or image)."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")
    if signature_type not in ["text", "image"]:
        raise HTTPException(status_code=400, detail="Signature type must be 'text' or 'image'")
    if signature_type == "text" and not signature_text:
        raise HTTPException(status_code=400, detail="Signature text is required")
    if signature_type == "image" and not signature_image:
        raise HTTPException(status_code=400, detail="Signature image is required")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 50MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)
        page_indices = _parse_pages(pages, total) if pages else list(range(total))

        position_map = {
            "top-left": (0.05, 0.07),
            "top-right": (0.60, 0.07),
            "bottom-left": (0.05, 0.95),
            "bottom-right": (0.60, 0.95),
            "center": (0.35, 0.50)
        }
        pos_x, pos_y = position_map.get(position, position_map["bottom-right"])

        if signature_type == "text":
            for idx in page_indices:
                if 0 <= idx < total:
                    page = doc[idx]
                    rect = page.rect
                    x = rect.width * pos_x
                    y = rect.height * pos_y
                    page.insert_text(
                        (x, y), signature_text,
                        fontsize=font_size,
                        color=(0, 0, 0.4),
                        overlay=True
                    )

        elif signature_type == "image":
            sig_img_content = await signature_image.read()
            if len(sig_img_content) == 0:
                raise HTTPException(status_code=400, detail="Empty signature image")
            if len(sig_img_content) > 5 * 1024 * 1024:
                raise HTTPException(status_code=413, detail="Signature image too large. Maximum 5MB")

            for idx in page_indices:
                if 0 <= idx < total:
                    page = doc[idx]
                    rect = page.rect
                    sig_width = rect.width * 0.15
                    sig_height = sig_width * 0.3
                    x = rect.width * pos_x
                    y = rect.height * pos_y
                    sig_rect = fitz.Rect(x, y, x + sig_width, y + sig_height)
                    page.insert_image(sig_rect, stream=sig_img_content, keep_proportion=True, overlay=True)

        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)

        return StreamingResponse(output, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=signed.pdf"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Signature failed: {str(e)}")


# ========== PHASE 3: SPLIT BY TEXT PATTERN ==========
@router.post("/split-pattern")
async def split_by_pattern(
    file: UploadFile = File(...),
    pattern: str = Form(...),
    mode: str = Form("exact")
):
    """Split PDF when text pattern is found on a page. Mode: exact, contains, regex."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")
    if not pattern.strip():
        raise HTTPException(status_code=400, detail="Pattern cannot be empty")
    if mode not in ["exact", "contains", "regex"]:
        raise HTTPException(status_code=400, detail="Mode must be 'exact', 'contains', or 'regex'")

    try:
        import re
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 100MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)

        split_pages = []
        for i in range(total):
            text = doc[i].get_text()
            match_found = False
            if mode == "exact":
                match_found = pattern in text
            elif mode == "contains":
                match_found = pattern.lower() in text.lower()
            elif mode == "regex":
                try:
                    match_found = bool(re.search(pattern, text, re.IGNORECASE))
                except re.error:
                    raise HTTPException(status_code=400, detail="Invalid regex pattern")

            if match_found:
                split_pages.append(i)

        if not split_pages:
            raise HTTPException(status_code=400, detail="Pattern not found in any page")

        sections = []
        for idx, split_idx in enumerate(split_pages):
            end = split_pages[idx + 1] if idx + 1 < len(split_pages) else total
            sections.append({"start": split_idx, "end": end})

        output_files = []
        for idx, sec in enumerate(sections):
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=sec["start"], to_page=sec["end"] - 1)
            buf = io.BytesIO()
            new_doc.save(buf)
            new_doc.close()
            buf.seek(0)
            output_files.append((f"section_{idx + 1}.pdf", buf))

        doc.close()

        if len(output_files) == 1:
            return StreamingResponse(output_files[0][1], media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"})

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, buf in output_files:
                zf.writestr(fname, buf.getvalue())
        zip_buf.seek(0)

        return StreamingResponse(zip_buf, media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=pattern_split.zip"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Pattern split failed: {str(e)}")


# ========== PHASE 3: SPLIT BY INVOICE NUMBER ==========
@router.post("/split-invoice")
async def split_by_invoice(
    file: UploadFile = File(...),
    pattern: str = Form(r"(?:invoice|order|po|ref)\s*(?:nr\.?|no\.?|number)?[:\s#]*(\d+)"),
    dpi: int = Form(150)
):
    """Split PDF by invoice/order numbers detected via OCR."""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF files allowed")

    try:
        import re
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
        ocr = RapidOCR()
    except ImportError:
        raise HTTPException(status_code=501, detail="OCR engine not available")

    try:
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty PDF file")
        if len(content) > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum 100MB")

        doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
        total = len(doc)
        mat = fitz.Matrix(dpi / 72, dpi / 72)

        invoices = []
        for i in range(total):
            text = doc[i].get_text()
            if not text.strip():
                pix = doc[i].get_pixmap(matrix=mat)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
                result, _ = ocr(img)
                text = " ".join([r[1] for r in result]) if result else ""

            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                invoice_num = match.group(1) if match.lastindex else match.group(0)
                invoices.append((i, invoice_num.strip()))

        if not invoices:
            raise HTTPException(status_code=400, detail="No invoice/order numbers found")

        sections = []
        for idx, (page_idx, inv_num) in enumerate(invoices):
            end = invoices[idx + 1][0] if idx + 1 < len(invoices) else total
            sections.append({"start": page_idx, "end": end, "invoice": inv_num})

        output_files = []
        for sec in sections:
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=sec["start"], to_page=sec["end"] - 1)
            buf = io.BytesIO()
            new_doc.save(buf)
            new_doc.close()
            buf.seek(0)
            safe_inv = "".join(c if c.isalnum() or c in "-_" else "_" for c in sec["invoice"])[:30]
            fname = f"invoice_{safe_inv}.pdf"
            output_files.append((fname, buf))

        doc.close()

        if len(output_files) == 1:
            return StreamingResponse(output_files[0][1], media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={output_files[0][0]}"})

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, buf in output_files:
                zf.writestr(fname, buf.getvalue())
        zip_buf.seek(0)

        return StreamingResponse(zip_buf, media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=invoice_split.zip"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Invoice split failed: {str(e)}")


# ========== HELPER ==========
def _parse_pages(pages_str: str, total: int) -> list:
    """Parse page specification string like '1-3, 5, 7-9' into 0-based indices."""
    indices = []
    for part in pages_str.split(","):
        part = part.strip()
        if "-" in part:
            start, end = map(int, part.split("-"))
            indices.extend(range(max(1, start) - 1, min(total, end)))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.append(idx)
    return indices


# ========== UTILITIES ==========
@router.get("/tools")
def list_tools():
    """List all available tools."""
    return {
        "tools": [
            {"slug": "merge-pdf", "endpoint": "/api/merge", "method": "POST", "phase": 1},
            {"slug": "split-pdf", "endpoint": "/api/split", "method": "POST", "phase": 1},
            {"slug": "compress-pdf", "endpoint": "/api/compress", "method": "POST", "phase": 1},
            {"slug": "pdf-to-word", "endpoint": "/api/pdf-to-word", "method": "POST", "phase": 1},
            {"slug": "word-to-pdf", "endpoint": "/api/word-to-pdf", "method": "POST", "phase": 1},
            {"slug": "jpg-to-pdf", "endpoint": "/api/jpg-to-pdf", "method": "POST", "phase": 1},
            {"slug": "pdf-to-jpg", "endpoint": "/api/pdf-to-jpg", "method": "POST", "phase": 1},
            {"slug": "rotate-pdf", "endpoint": "/api/rotate", "method": "POST", "phase": 2},
            {"slug": "delete-pages", "endpoint": "/api/delete-pages", "method": "POST", "phase": 2},
            {"slug": "extract-pages", "endpoint": "/api/extract-pages", "method": "POST", "phase": 2},
            {"slug": "watermark-pdf", "endpoint": "/api/watermark", "method": "POST", "phase": 2},
            {"slug": "protect-pdf", "endpoint": "/api/protect", "method": "POST", "phase": 2},
            {"slug": "unlock-pdf", "endpoint": "/api/unlock", "method": "POST", "phase": 2},
            {"slug": "sign-pdf", "endpoint": "/api/sign-pdf", "method": "POST", "phase": 2},
            {"slug": "split-pattern", "endpoint": "/api/split-pattern", "method": "POST", "phase": 3},
            {"slug": "split-invoice", "endpoint": "/api/split-invoice", "method": "POST", "phase": 3},
            {"slug": "split-barcode", "endpoint": "/api/split-barcode", "method": "POST", "phase": 3},
            {"slug": "auto-rename", "endpoint": "/api/auto-rename", "method": "POST", "phase": 3},
            {"slug": "detect-type", "endpoint": "/api/detect-type", "method": "POST", "phase": 3},
            {"slug": "process-archive", "endpoint": "/api/process-archive", "method": "POST", "phase": 4}
        ]
    }
